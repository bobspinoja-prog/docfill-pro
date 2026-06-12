import re
from pathlib import Path
from typing import Any

from docx import Document


PLACEHOLDER_PATTERN = re.compile(r"\{\{[^{}]+\}\}")


class DOCXReader:
    """Lê um documento Word e gera texto para preview."""

    FIELD_MARKERS = {
        "comprador": "{{COMPRADOR}}",
        "nacionalidade": "{{NACIONALIDADE}}",
        "profissao": "{{PROFISSAO}}",
        "estado_civil": "{{ESTADO_CIVIL}}",
        "cpf_cnpj": "{{CPF_CNPJ}}",
        "lote": "{{LOTE}}",
        "quadra": "{{QUADRA}}",
        "empreendimento": "{{EMPREENDIMENTO}}",
        "vendedor": "{{VENDEDOR}}",
        "cidade": "{{CIDADE}}",
        "data": "{{DATA}}",
    }

    def __init__(self, template_path: str | Path) -> None:
        self.template_path = Path(template_path)
        self._document: Any | None = None
        self._preview_blocks: list[str] | None = None
        self._analysis: dict[str, Any] | None = None

    def extract_text(self, replacements: dict | None = None) -> str:
        replacements = replacements or {}
        blocks = [
            self._replace_text(block, replacements)
            for block in self._get_preview_blocks()
        ]
        visible_blocks = [block for block in blocks if block.strip()]
        return "\n\n".join(visible_blocks)

    def analyze_template(self) -> dict[str, Any]:
        """Analisa um template e retorna marcadores e áreas encontradas."""
        if self._analysis is None:
            self._analysis = self._build_analysis()
        return self._copy_analysis(self._analysis)

    def suggest_values(self) -> dict[str, str]:
        """Extrai valores editáveis de documentos já preenchidos."""
        text = "\n".join(self._get_preview_blocks())
        suggestions: dict[str, str] = {}

        declaration_pattern = re.compile(
            r"(?P<comprador>[^,\n]+),\s*"
            r"(?P<nacionalidade>[^,\n]+),\s*"
            r"(?P<profissao>[^,\n]+),\s*"
            r"(?P<estado_civil>[^,\n]+),\s*"
            r"portador(?:a)?\s+do\s+CPF\s+n[°º]\s*"
            r"(?P<cpf_cnpj>.*?)\s+na\s+qualidade\s+de\s+COMPRADOR\s+do\s+"
            r"Lote\s+(?P<lote>.*?)\s+Quadra\s+(?P<quadra>.*?)\s+do\s+"
            r"(?P<empreendimento>.*?),\s+adquirido.*?do\(a\)\s+"
            r"(?P<vendedor>.*?),\s+para",
            re.IGNORECASE | re.DOTALL,
        )
        match = declaration_pattern.search(text)
        if match:
            for group_name, marker in self.FIELD_MARKERS.items():
                if group_name in match.groupdict():
                    value = self._normalize_spaces(match.group(group_name))
                    if value:
                        suggestions[marker] = value

        date_pattern = re.compile(
            r"(?P<cidade>[A-Za-zÀ-ÿ][A-Za-zÀ-ÿ '-]{1,80}),\s*"
            r"(?P<data>\d{1,2}\s+de\s+[A-Za-zÀ-ÿ]+\s+de\s+\d{4})",
            re.IGNORECASE,
        )
        date_match = date_pattern.search(text)
        if date_match:
            suggestions.setdefault("{{CIDADE}}", self._normalize_spaces(date_match.group("cidade")))
            suggestions.setdefault("{{DATA}}", self._normalize_spaces(date_match.group("data")))

        return suggestions

    def build_literal_replacements(self, values: dict[str, str]) -> dict[str, str]:
        """Cria substituições para documentos sem marcadores, usando valores detectados."""
        suggestions = self.suggest_values()
        literal_replacements: dict[str, str] = {}

        direct_markers = (
            "{{COMPRADOR}}",
            "{{NACIONALIDADE}}",
            "{{PROFISSAO}}",
            "{{ESTADO_CIVIL}}",
            "{{CPF_CNPJ}}",
            "{{EMPREENDIMENTO}}",
            "{{VENDEDOR}}",
        )
        for marker in direct_markers:
            source = suggestions.get(marker, "").strip()
            target = values.get(marker, "").strip()
            if source and target and source != target:
                literal_replacements[source] = target

        self._add_prefixed_literal(
            literal_replacements,
            "Lote",
            suggestions.get("{{LOTE}}", ""),
            values.get("{{LOTE}}", ""),
        )
        self._add_prefixed_literal(
            literal_replacements,
            "Quadra",
            suggestions.get("{{QUADRA}}", ""),
            values.get("{{QUADRA}}", ""),
        )

        current_city = suggestions.get("{{CIDADE}}", "").strip()
        current_date = suggestions.get("{{DATA}}", "").strip()
        new_city = values.get("{{CIDADE}}", "").strip() or current_city
        new_date = values.get("{{DATA}}", "").strip() or current_date
        if current_city and current_date and (new_city != current_city or new_date != current_date):
            literal_replacements[f"{current_city}, {current_date}"] = f"{new_city}, {new_date}".strip(", ")

        return literal_replacements

    def _get_document(self) -> Any:
        if not self.template_path.exists():
            raise FileNotFoundError("Arquivo template não encontrado.")
        if self._document is None:
            self._document = Document(self.template_path)
        return self._document

    def _get_preview_blocks(self) -> list[str]:
        if self._preview_blocks is None:
            self._preview_blocks = self._collect_preview_blocks()
        return self._preview_blocks

    def _collect_preview_blocks(self) -> list[str]:
        document = self._get_document()
        blocks: list[str] = []

        for paragraph in document.paragraphs:
            self._append_block(blocks, paragraph.text)

        for table in document.tables:
            blocks.extend(self._extract_table_rows(table))

        for label, part in self._iter_section_parts(document):
            for paragraph in part.paragraphs:
                self._append_block(blocks, paragraph.text, label)
            for table in part.tables:
                for row_text in self._extract_table_rows(table):
                    self._append_block(blocks, row_text, label)

        return blocks

    def _build_analysis(self) -> dict[str, Any]:
        document = self._get_document()
        analysis: dict[str, Any] = {
            "template": self.template_path.name,
            "paragraphs": 0,
            "tables": 0,
            "headers": 0,
            "footers": 0,
            "placeholders": [],
            "areas": [],
            "summary": [],
        }

        for paragraph in document.paragraphs:
            analysis["paragraphs"] += 1
            self._track_placeholders(analysis, paragraph.text, "Parágrafos")

        for table in self._iter_tables(document.tables):
            analysis["tables"] += 1
            for paragraph_text in self._iter_table_paragraph_text(table):
                self._track_placeholders(analysis, paragraph_text, "Tabelas")

        for label, part in self._iter_section_parts(document):
            has_content = any(paragraph.text.strip() for paragraph in part.paragraphs) or bool(part.tables)
            if not has_content:
                continue

            if label == "Cabeçalho":
                analysis["headers"] += 1
                area = "Cabeçalhos"
            else:
                analysis["footers"] += 1
                area = "Rodapés"

            for paragraph in part.paragraphs:
                self._track_placeholders(analysis, paragraph.text, area)

            for table in self._iter_tables(part.tables):
                for paragraph_text in self._iter_table_paragraph_text(table):
                    self._track_placeholders(analysis, paragraph_text, area)

        unique_markers = list(dict.fromkeys(analysis["placeholders"]))
        analysis["placeholders"] = unique_markers
        analysis["areas"] = list(dict.fromkeys(analysis["areas"]))
        analysis["summary"] = [
            f"Parágrafos: {analysis['paragraphs']}",
            f"Tabelas: {analysis['tables']}",
            f"Cabeçalhos: {analysis['headers']}",
            f"Rodapés: {analysis['footers']}",
            f"Marcadores únicos: {len(unique_markers)}",
        ]
        return analysis

    @staticmethod
    def _append_block(blocks: list[str], text: str, label: str | None = None) -> None:
        clean_text = text.strip()
        if not clean_text:
            return
        blocks.append(f"[{label}] {clean_text}" if label else clean_text)

    @classmethod
    def _extract_table_rows(cls, table: Any) -> list[str]:
        rows: list[str] = []
        for row in table.rows:
            values: list[str] = []
            for cell in row.cells:
                cell_text = "\n".join(
                    paragraph.text.strip()
                    for paragraph in cell.paragraphs
                    if paragraph.text.strip()
                )
                if cell_text:
                    values.append(cell_text)
                for nested_table in cell.tables:
                    nested_text = " / ".join(cls._extract_table_rows(nested_table))
                    if nested_text:
                        values.append(nested_text)
            if values:
                rows.append(" | ".join(values))
        return rows

    @classmethod
    def _iter_tables(cls, tables: Any) -> Any:
        for table in tables:
            yield table
            for row in table.rows:
                for cell in row.cells:
                    yield from cls._iter_tables(cell.tables)

    @classmethod
    def _iter_table_paragraph_text(cls, table: Any) -> Any:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph.text
                for nested_table in cell.tables:
                    yield from cls._iter_table_paragraph_text(nested_table)

    @staticmethod
    def _iter_section_parts(document: Any) -> Any:
        seen: set[int] = set()
        part_names = (
            ("Cabeçalho", "header"),
            ("Cabeçalho", "first_page_header"),
            ("Cabeçalho", "even_page_header"),
            ("Rodapé", "footer"),
            ("Rodapé", "first_page_footer"),
            ("Rodapé", "even_page_footer"),
        )
        for section in document.sections:
            for label, attr_name in part_names:
                part = getattr(section, attr_name, None)
                if part is None:
                    continue
                element_id = id(part._element)
                if element_id in seen:
                    continue
                seen.add(element_id)
                yield label, part

    @staticmethod
    def _track_placeholders(analysis: dict[str, Any], text: str, area: str) -> None:
        found = PLACEHOLDER_PATTERN.findall(text or "")
        if found:
            analysis["placeholders"].extend(found)
            analysis["areas"].append(area)

    @staticmethod
    def _normalize_spaces(value: str) -> str:
        return " ".join((value or "").split())

    @staticmethod
    def _add_prefixed_literal(replacements: dict[str, str], prefix: str, source: str, target: str) -> None:
        clean_source = DOCXReader._normalize_spaces(source)
        clean_target = DOCXReader._normalize_spaces(target)
        if not clean_source or not clean_target or clean_source == clean_target:
            return
        if clean_target.lower().startswith(prefix.lower()):
            replacements[f"{prefix} {clean_source}"] = clean_target
        else:
            replacements[f"{prefix} {clean_source}"] = f"{prefix} {clean_target}"

    @staticmethod
    def _copy_analysis(analysis: dict[str, Any]) -> dict[str, Any]:
        return {
            **analysis,
            "placeholders": list(analysis["placeholders"]),
            "areas": list(analysis["areas"]),
            "summary": list(analysis["summary"]),
        }

    @staticmethod
    def _replace_text(text: str, replacements: dict) -> str:
        if not text or not replacements:
            return text

        result = text
        for marker, value in replacements.items():
            marker_text = str(marker)
            if not marker_text:
                continue
            replacement = "" if value is None else str(value)
            result = result.replace(marker_text, replacement)
        return result
