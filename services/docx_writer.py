from pathlib import Path
from typing import Any

from docx import Document


class DOCXWriter:
    """Gera um novo documento Word preenchido com os dados informados."""

    def generate(self, template_path: str | Path, output_path: str | Path, replacements: dict) -> Path:
        template = Path(template_path)
        output = Path(output_path)
        clean_replacements = self._clean_replacements(replacements)

        if not template.exists():
            raise FileNotFoundError("Template não encontrado.")
        if template.resolve() == output.resolve():
            raise ValueError("O arquivo gerado não pode sobrescrever o template original.")

        output.parent.mkdir(parents=True, exist_ok=True)

        document = Document(template)
        self._replace_in_paragraphs(document.paragraphs, clean_replacements)
        self._replace_in_tables(document.tables, clean_replacements)
        self._replace_in_sections(document, clean_replacements)

        document.save(output)
        return output

    @staticmethod
    def _replace_in_paragraphs(paragraphs: Any, replacements: dict[str, str]) -> None:
        for paragraph in paragraphs:
            DOCXWriter._apply_replacements(paragraph, replacements)

    @classmethod
    def _replace_in_tables(cls, tables: Any, replacements: dict[str, str]) -> None:
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        cls._apply_replacements(paragraph, replacements)
                    cls._replace_in_tables(cell.tables, replacements)

    @classmethod
    def _replace_in_sections(cls, document: Any, replacements: dict[str, str]) -> None:
        for part in cls._iter_section_parts(document):
            cls._replace_in_paragraphs(part.paragraphs, replacements)
            cls._replace_in_tables(part.tables, replacements)

    @classmethod
    def _apply_replacements(cls, paragraph: Any, replacements: dict[str, str]) -> None:
        if not replacements or not paragraph.text:
            return

        full_text = "".join(run.text for run in paragraph.runs)
        replacement_count = sum(full_text.count(marker) for marker in replacements if marker)
        if replacement_count == 0:
            return

        for _ in range(replacement_count):
            full_text = "".join(run.text for run in paragraph.runs)
            match = cls._find_next_marker(full_text, replacements)
            if match is None:
                return

            start, marker = match
            end = start + len(marker)
            cls._replace_text_range(paragraph.runs, start, end, replacements[marker])

    @staticmethod
    def _replace_text(text: str, replacements: dict[str, str]) -> str:
        result = text
        for marker, value in replacements.items():
            result = result.replace(marker, value)
        return result

    @staticmethod
    def _find_next_marker(text: str, replacements: dict[str, str]) -> tuple[int, str] | None:
        matches = [
            (index, marker)
            for marker in replacements
            if marker and (index := text.find(marker)) >= 0
        ]
        if not matches:
            return None
        return min(matches, key=lambda item: item[0])

    @staticmethod
    def _replace_text_range(runs: Any, start: int, end: int, replacement: str) -> None:
        positions: list[tuple[int, int]] = []
        for run_index, run in enumerate(runs):
            positions.extend((run_index, char_index) for char_index, _ in enumerate(run.text))

        if not positions or start >= len(positions) or end <= start:
            return

        start_run_index, start_char_index = positions[start]
        end_run_index, end_char_index = positions[end - 1]
        start_run = runs[start_run_index]
        end_run = runs[end_run_index]

        before = start_run.text[:start_char_index]
        after = end_run.text[end_char_index + 1:]

        if start_run_index == end_run_index:
            start_run.text = f"{before}{replacement}{after}"
            return

        start_run.text = f"{before}{replacement}"
        for run_index in range(start_run_index + 1, end_run_index):
            runs[run_index].text = ""
        end_run.text = after

    @staticmethod
    def _clean_replacements(replacements: dict) -> dict[str, str]:
        clean: dict[str, str] = {}
        for marker, value in replacements.items():
            marker_text = str(marker)
            if not marker_text:
                continue
            clean[marker_text] = "" if value is None else str(value)
        return clean

    @staticmethod
    def _iter_section_parts(document: Any) -> Any:
        seen: set[int] = set()
        part_names = (
            "header",
            "first_page_header",
            "even_page_header",
            "footer",
            "first_page_footer",
            "even_page_footer",
        )
        for section in document.sections:
            for attr_name in part_names:
                part = getattr(section, attr_name, None)
                if part is None:
                    continue
                element_id = id(part._element)
                if element_id in seen:
                    continue
                seen.add(element_id)
                yield part
