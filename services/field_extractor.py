from __future__ import annotations

import re
import unicodedata
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

from docx import Document

from services.docx_writer import DOCXWriter
from services.text_sections import DocumentSections, split_text_sections


FIELDS = (
    "COMPRADOR",
    "NACIONALIDADE",
    "PROFISSAO",
    "ESTADO_CIVIL",
    "CPF_CNPJ",
    "LOTE",
    "QUADRA",
    "EMPREENDIMENTO",
    "VENDEDOR",
    "CIDADE",
    "DATA",
)

MARKER_BY_FIELD = {field_name: f"{{{{{field_name}}}}}" for field_name in FIELDS}
FIELD_BY_MARKER = {marker: field_name for field_name, marker in MARKER_BY_FIELD.items()}

LETTER_RANGE = r"A-Za-z\u00C0-\u017F\?"
UPPER_RANGE = r"A-Z\u00C0-\u00D6\u00D8-\u00DE\?"
UPPER_WORD = rf"[{UPPER_RANGE}0-9&.'\-]+"
PERSON_NAME_PATTERN = rf"[{UPPER_RANGE}][{UPPER_RANGE}0-9&.'\-]+(?:\s+{UPPER_WORD}){{1,12}}"
CPF_CNPJ_PATTERN = r"(?:\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}|\d{3}\.\d{3}\.\d{3}-\d{2})"
LOOSE_ID_PATTERN = r"(?:\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}|\d{3}\.\d{3}\.\d{3}-\d{2}|[\d./-]{11,20})"
DATE_PATTERN = rf"\d{{1,2}}\s+de\s+[{LETTER_RANGE}]+\s+de\s+\d{{4}}"
MARITAL_PATTERN = (
    r"(?i:solteir[oa]|casad[oa]|divorciad[oa]|vi[\u00FAu]v[oa]|separad[oa]|"
    r"desquitad[oa]|uni[\u00E3a]o\s+est[\u00E1a]vel|convivente)"
)
NATIONALITY_WORDS = {
    "brasileiro",
    "brasileira",
    "estrangeiro",
    "estrangeira",
    "portugues",
    "portuguesa",
    "portugu\u00EAs",
    "italiano",
    "italiana",
    "espanhol",
    "espanhola",
    "argentino",
    "argentina",
    "alemao",
    "alema",
    "alem\u00E3o",
    "alem\u00E3",
}


@dataclass
class FieldOccurrence:
    value: str
    start_index: int
    end_index: int
    section: str
    replace: bool = True
    reason: str = ""

    def to_dict(self) -> dict[str, str | int | bool]:
        return asdict(self)


@dataclass
class ExtractedField:
    value: str = ""
    confidence: float = 0.0
    source: str = "not_found"
    reason: str = ""
    start_index: int = -1
    end_index: int = -1
    section: str = ""
    marker: str = ""
    evidence: str = ""
    occurrences: list[FieldOccurrence] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        data = asdict(self)
        data["occurrences"] = [occurrence.to_dict() for occurrence in self.occurrences]
        return data


@dataclass
class FieldExtractionResult:
    fields: dict[str, ExtractedField]
    normalized_text: str
    markers: list[str]
    sections: DocumentSections

    def to_dict(self) -> dict[str, dict[str, Any]]:
        return {field_name: self.fields[field_name].to_dict() for field_name in FIELDS}

    def as_plain_values(self, min_confidence: float = 0.0) -> dict[str, str]:
        return {
            field_name: item.value
            for field_name, item in self.fields.items()
            if item.value and item.confidence >= min_confidence
        }

    def as_marker_values(self, min_confidence: float = 0.0) -> dict[str, str]:
        return {
            MARKER_BY_FIELD[field_name]: item.value
            for field_name, item in self.fields.items()
            if item.value and item.confidence >= min_confidence
        }


@dataclass
class TemplateRewriteReport:
    output_path: Path
    marked_fields: list[str] = field(default_factory=list)
    ignored_fields: list[str] = field(default_factory=list)
    review_fields: list[str] = field(default_factory=list)
    replacements: dict[str, str] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return {
            "output_path": str(self.output_path),
            "marked_fields": list(dict.fromkeys(self.marked_fields)),
            "ignored_fields": list(dict.fromkeys(self.ignored_fields)),
            "review_fields": list(dict.fromkeys(self.review_fields)),
            "replacements": dict(self.replacements),
        }


def normalize_text(text: str) -> str:
    normalized, _mapping = _normalize_text_with_mapping(text)
    return normalized


def detect_markers(text: str) -> list[str]:
    markers = re.findall(r"\{\{[^{}]+\}\}", text or "")
    return list(dict.fromkeys(markers))


def extract_fields(text: str) -> FieldExtractionResult:
    normalized, index_map = _normalize_text_with_mapping(text)
    sections = split_text_sections(normalized)
    fields = {
        field_name: ExtractedField(marker=MARKER_BY_FIELD[field_name])
        for field_name in FIELDS
    }

    _extract_preamble(sections, index_map, normalized, fields)
    _extract_property(sections, index_map, normalized, fields)
    _extract_seller(sections, index_map, normalized, fields)
    _extract_city_date(sections, index_map, normalized, fields)
    _cross_validate(sections, index_map, normalized, fields)

    return FieldExtractionResult(
        fields=fields,
        normalized_text=normalized,
        markers=detect_markers(text),
        sections=sections,
    )


def rewrite_template_with_markers(
    input_docx: str | Path,
    output_docx: str | Path,
    extraction_result: FieldExtractionResult,
) -> TemplateRewriteReport:
    input_path = Path(input_docx)
    output_path = Path(output_docx)
    if not input_path.exists():
        raise FileNotFoundError("Template original nao encontrado.")
    if input_path.resolve() == output_path.resolve():
        raise ValueError("O template marcado nao pode sobrescrever o original.")

    document_text = _collect_docx_text(input_path)
    report = TemplateRewriteReport(output_path=output_path)
    candidates: dict[str, str] = {}
    expected_counts: dict[str, int] = {}
    field_by_value: dict[str, set[str]] = {}

    for field_name, item in extraction_result.fields.items():
        if not item.value:
            report.ignored_fields.append(f"{field_name}: sem valor detectado")
            continue
        if item.confidence < 0.60:
            report.ignored_fields.append(f"{field_name}: confianca abaixo de 0.60")
            continue
        if item.confidence < 0.85:
            report.review_fields.append(f"{field_name}: revisar antes de marcar ({item.confidence:.2f})")
            continue

        replace_occurrences = [occurrence for occurrence in item.occurrences if occurrence.replace and occurrence.value]
        if not replace_occurrences:
            report.ignored_fields.append(f"{field_name}: sem ocorrencia segura")
            continue

        for occurrence in replace_occurrences:
            candidates[occurrence.value] = item.marker
            expected_counts[occurrence.value] = expected_counts.get(occurrence.value, 0) + 1
            field_by_value.setdefault(occurrence.value, set()).add(field_name)

    safe_replacements: dict[str, str] = {}
    for source_value, marker in sorted(candidates.items(), key=lambda item: len(item[0]), reverse=True):
        actual_count = document_text.count(source_value)
        expected_count = expected_counts.get(source_value, 0)
        fields_for_value = ", ".join(sorted(field_by_value.get(source_value, set())))
        if actual_count == 0:
            report.ignored_fields.append(f"{fields_for_value}: valor nao encontrado no DOCX")
            continue
        if actual_count != expected_count:
            report.ignored_fields.append(
                f"{fields_for_value}: ocorrencia ambigua de '{source_value}' ({actual_count} no DOCX, {expected_count} confiavel)"
            )
            continue
        safe_replacements[source_value] = marker
        report.marked_fields.extend(sorted(field_by_value.get(source_value, set())))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document(input_path)
    _replace_in_document(document, safe_replacements)
    document.save(output_path)
    report.replacements = safe_replacements
    return report


def _extract_preamble(
    sections: DocumentSections,
    index_map: list[int],
    full_text: str,
    fields: dict[str, ExtractedField],
) -> None:
    section = sections.get("preambulo")
    text = section.text
    pattern = re.compile(
        rf"^\s*(?P<comprador>{PERSON_NAME_PATTERN})\s*,\s*"
        rf"(?P<qualificacao>.*?)"
        rf"\s*,?\s*(?i:portador(?:a)?|inscrit[oa])\s+(?i:do|da|no|na)?\s*"
        rf"(?i:CPF|CNPJ|CPF/CNPJ)\s*n?[\u00BA\u00B0o.\?]?\s*(?P<cpf>{LOOSE_ID_PATTERN})",
        re.DOTALL,
    )
    match = pattern.search(text)
    if not match:
        _extract_document_id(section, index_map, full_text, fields)
        return

    comprador_span = _global_span(section, match.start("comprador"), match.end("comprador"))
    _set_field(
        fields,
        "COMPRADOR",
        match.group("comprador"),
        0.85,
        "preambulo",
        "nome em caixa alta no inicio do preambulo",
        section.name,
        _original_span(index_map, *comprador_span),
        _snippet(full_text, *comprador_span),
    )

    parts = _comma_parts_with_spans(match.group("qualificacao"), section.start_index + match.start("qualificacao"))
    nationality = _find_nationality(parts)
    marital_status = _find_marital_status(parts)

    if nationality:
        _set_field(
            fields,
            "NACIONALIDADE",
            nationality["value"],
            0.95,
            "preambulo",
            "primeiro termo apos comprador",
            section.name,
            _original_span(index_map, nationality["start"], nationality["end"]),
            _snippet(full_text, nationality["start"], nationality["end"]),
        )

    if marital_status:
        _set_field(
            fields,
            "ESTADO_CIVIL",
            marital_status["value"],
            0.95,
            "preambulo",
            "estado civil antes do CPF",
            section.name,
            _original_span(index_map, marital_status["start"], marital_status["end"]),
            _snippet(full_text, marital_status["start"], marital_status["end"]),
        )

    professions = [
        part
        for part in parts
        if _lower_ascii(part["value"]) not in {
            _lower_ascii(nationality["value"]) if nationality else "",
            _lower_ascii(marital_status["value"]) if marital_status else "",
        }
        and not _is_rejected_profession_part(part["value"])
    ]
    if professions:
        profession_occurrences = [
            FieldOccurrence(
                value=part["value"],
                start_index=_original_span(index_map, part["start"], part["end"])[0],
                end_index=_original_span(index_map, part["start"], part["end"])[1],
                section=section.name,
                replace=True,
                reason="profissao antes do CPF",
            )
            for part in professions
        ]
        _set_field(
            fields,
            "PROFISSAO",
            " / ".join(dict.fromkeys(part["value"] for part in professions)),
            0.90,
            "preambulo",
            "termos restantes entre nacionalidade/estado civil e CPF",
            section.name,
            (profession_occurrences[0].start_index, profession_occurrences[-1].end_index),
            _snippet(full_text, professions[0]["start"], professions[-1]["end"]),
            occurrences=profession_occurrences,
        )

    cpf_span = _global_span(section, match.start("cpf"), match.end("cpf"))
    _set_field(
        fields,
        "CPF_CNPJ",
        match.group("cpf"),
        0.98,
        "preambulo",
        "CPF/CNPJ proximo do rotulo CPF/CNPJ",
        section.name,
        _original_span(index_map, *cpf_span),
        _snippet(full_text, *cpf_span),
    )


def _extract_document_id(
    section: Any,
    index_map: list[int],
    full_text: str,
    fields: dict[str, ExtractedField],
) -> None:
    text = section.text if hasattr(section, "text") else full_text
    start_offset = section.start_index if hasattr(section, "start_index") else 0
    context_match = re.search(
        rf"(?i:(?:CPF|CNPJ|CPF/CNPJ))\s*n?[\u00BA\u00B0o.\?]?\s*(?P<cpf>{LOOSE_ID_PATTERN})",
        text,
    )
    if context_match:
        span = (start_offset + context_match.start("cpf"), start_offset + context_match.end("cpf"))
        _set_field(fields, "CPF_CNPJ", context_match.group("cpf"), 0.98, "preambulo", "CPF/CNPJ por contexto", "preambulo", _original_span(index_map, *span), _snippet(full_text, *span))
        return
    loose_match = re.search(LOOSE_ID_PATTERN, text)
    if loose_match:
        span = (start_offset + loose_match.start(), start_offset + loose_match.end())
        _set_field(fields, "CPF_CNPJ", loose_match.group(0), 0.90, "fallback", "padrao de CPF/CNPJ sem rotulo", "preambulo", _original_span(index_map, *span), _snippet(full_text, *span))


def _extract_property(
    sections: DocumentSections,
    index_map: list[int],
    full_text: str,
    fields: dict[str, ExtractedField],
) -> None:
    section = sections.get("preambulo")
    pattern = re.compile(
        rf"(?i:na\s+qualidade\s+de\s+COMPRADOR\s+do\s+(?:Lote|Unidade)\s+)"
        rf"(?P<lote>[A-Z0-9][A-Z0-9./\-]*)"
        rf"(?:\s+(?i:Quadra)\s+(?P<quadra>[A-Z0-9][A-Z0-9./\-]*))?"
        rf"\s+(?i:do|da)\s+"
        rf"(?P<empreendimento>(?i:LOTEAMENTO|CONDOM[\u00CDI]NIO|CONDOMINIO|EMPREENDIMENTO|RESIDENCIAL|EDIF[\u00CDI]CIO|ALPHAVILLE)"
        rf"[{UPPER_RANGE}0-9 .'/\-]{{4,160}}?)"
        rf"(?=,|\.|\s+(?i:adquirido|localizado|situado)|$)",
        re.DOTALL,
    )
    match = pattern.search(section.text)
    if not match:
        return

    for field_name, group_name, reason in (
        ("LOTE", "lote", "codigo apos Lote/Unidade do comprador"),
        ("QUADRA", "quadra", "codigo imediatamente apos Quadra"),
        ("EMPREENDIMENTO", "empreendimento", "empreendimento apos lote e quadra"),
    ):
        if not match.groupdict().get(group_name):
            continue
        span = _global_span(section, match.start(group_name), match.end(group_name))
        _set_field(
            fields,
            field_name,
            match.group(group_name),
            0.95,
            "preambulo",
            reason,
            section.name,
            _original_span(index_map, *span),
            _snippet(full_text, *span),
        )


def _extract_seller(
    sections: DocumentSections,
    index_map: list[int],
    full_text: str,
    fields: dict[str, ExtractedField],
) -> None:
    preambulo = sections.get("preambulo")
    clause_patterns = (
        rf"(?i:adquirido\s+(?:atrav[\u00E9e]s\s+de\s+)?).{{0,220}}?"
        rf"(?i:do\(a\)|da|do)\s+(?i:Sr\.?\(a\)\.?\s*)?(?P<vendedor>{PERSON_NAME_PATTERN})",
        rf"(?i:(?:compra\s+e\s+venda|promessa\s+de\s+venda\s+e\s+compra|com\s+permuta)).{{0,120}}?"
        rf"(?i:do\(a\)|da|do)\s+(?i:Sr\.?\(a\)\.?\s*)?(?P<vendedor>{PERSON_NAME_PATTERN})",
    )
    for pattern in clause_patterns:
        match = re.search(pattern, preambulo.text, re.DOTALL)
        if not match:
            continue
        span = _global_span(preambulo, match.start("vendedor"), match.end("vendedor"))
        _set_field(
            fields,
            "VENDEDOR",
            match.group("vendedor"),
            0.97,
            "preambulo",
            "nome apos gatilho de aquisicao",
            preambulo.name,
            _original_span(index_map, *span),
            _snippet(full_text, *span),
        )
        return

    assinaturas = sections.get("assinaturas")
    signature = re.search(rf"(?P<vendedor>{PERSON_NAME_PATTERN})\s+(?i:VENDEDOR|CEDENTE|ALIENANTE)\b", assinaturas.text)
    if signature:
        span = _global_span(assinaturas, signature.start("vendedor"), signature.end("vendedor"))
        _set_field(fields, "VENDEDOR", signature.group("vendedor"), 0.72, "assinatura", "fallback pela assinatura esquerda", assinaturas.name, _original_span(index_map, *span), _snippet(full_text, *span))
        return

    final = sections.get("paragrafo_final")
    final_clause = re.search(
        rf"(?i:responsabilidade\s+sobre\s+o\s+meu\s+compromisso\s+firmado\s+com\s+o\(a\)\s+Sr\.?\(a\)\.?\s*)"
        rf"(?P<vendedor>{PERSON_NAME_PATTERN})",
        final.text,
        re.DOTALL,
    )
    if final_clause:
        span = _global_span(final, final_clause.start("vendedor"), final_clause.end("vendedor"))
        _set_field(fields, "VENDEDOR", final_clause.group("vendedor"), 0.55, "fallback", "paragrafo final com baixa confianca", final.name, _original_span(index_map, *span), _snippet(full_text, *span))


def _extract_city_date(
    sections: DocumentSections,
    index_map: list[int],
    full_text: str,
    fields: dict[str, ExtractedField],
) -> None:
    section = sections.get("data")
    if section.is_empty:
        return
    match = re.search(
        rf"(?P<cidade>[{LETTER_RANGE}][{LETTER_RANGE} .'\-]{{2,80}})\s*,\s*(?P<data>{DATE_PATTERN})",
        section.text,
        flags=re.IGNORECASE,
    )
    if not match:
        return
    for field_name, group_name in (("CIDADE", "cidade"), ("DATA", "data")):
        span = _global_span(section, match.start(group_name), match.end(group_name))
        _set_field(
            fields,
            field_name,
            match.group(group_name),
            0.95,
            "data",
            "linha de cidade/data",
            section.name,
            _original_span(index_map, *span),
            _snippet(full_text, *span),
        )


def _cross_validate(
    sections: DocumentSections,
    index_map: list[int],
    full_text: str,
    fields: dict[str, ExtractedField],
) -> None:
    signatures = sections.get("assinaturas")
    for field_name in ("COMPRADOR", "VENDEDOR"):
        item = fields[field_name]
        if not item.value or signatures.is_empty:
            continue
        for match in re.finditer(re.escape(item.value), signatures.text, flags=re.IGNORECASE):
            span = _global_span(signatures, match.start(), match.end())
            item.occurrences.append(
                FieldOccurrence(
                    value=item.value,
                    start_index=_original_span(index_map, *span)[0],
                    end_index=_original_span(index_map, *span)[1],
                    section=signatures.name,
                    replace=True,
                    reason="confirmado em assinatura",
                )
            )
            item.confidence = round(min(0.99, max(item.confidence, item.confidence + 0.03)), 2)
            item.source = f"{item.source}+assinatura"
            break

    seller = fields["VENDEDOR"]
    final = sections.get("paragrafo_final")
    if seller.value and not final.is_empty:
        conflict = re.search(
            rf"(?i:responsabilidade\s+sobre\s+o\s+meu\s+compromisso\s+firmado\s+com\s+o\(a\)\s+Sr\.?\(a\)\.?\s*)"
            rf"(?P<nome>{PERSON_NAME_PATTERN})",
            final.text,
        )
        if conflict:
            span = _global_span(final, conflict.start("nome"), conflict.end("nome"))
            conflict_name = _clean_name(conflict.group("nome"))
            replace = _same_text(conflict_name, seller.value)
            seller.occurrences.append(
                FieldOccurrence(
                    value=conflict_name,
                    start_index=_original_span(index_map, *span)[0],
                    end_index=_original_span(index_map, *span)[1],
                    section=final.name,
                    replace=replace,
                    reason="nome em paragrafo final" if replace else "nome conflitante em paragrafo final",
                )
            )
            if not replace:
                seller.confidence = round(max(0.0, seller.confidence - 0.03), 2)
                seller.reason = f"{seller.reason}; conflito no paragrafo final"


def _set_field(
    fields: dict[str, ExtractedField],
    field_name: str,
    value: str,
    confidence: float,
    source: str,
    reason: str,
    section: str,
    original_span: tuple[int, int],
    evidence: str,
    occurrences: list[FieldOccurrence] | None = None,
) -> None:
    clean = _clean_name(value) if field_name in {"COMPRADOR", "VENDEDOR"} else _clean_value(value)
    if field_name == "EMPREENDIMENTO":
        clean = _clean_enterprise(clean)
    if field_name == "CIDADE":
        clean = _clean_city(clean)
    if not clean:
        return

    current = fields[field_name]
    if current.value and current.confidence >= confidence:
        return

    start_index, end_index = original_span
    if occurrences is None:
        occurrences = [
            FieldOccurrence(
                value=clean,
                start_index=start_index,
                end_index=end_index,
                section=section,
                replace=True,
                reason=reason,
            )
        ]
    current.value = clean
    current.confidence = round(min(1.0, max(0.0, confidence)), 2)
    current.source = source
    current.reason = reason
    current.start_index = start_index
    current.end_index = end_index
    current.section = section
    current.marker = MARKER_BY_FIELD[field_name]
    current.evidence = evidence
    current.occurrences = occurrences


def _normalize_text_with_mapping(text: str) -> tuple[str, list[int]]:
    value = (text or "").replace("\r\n", "\n").replace("\r", "\n").replace("\xa0", " ").replace("\u00AD", "")
    value = re.sub(rf"([{LETTER_RANGE}])-\s*\n\s*([{LETTER_RANGE}])", r"\1\2", value)
    chars: list[str] = []
    mapping: list[int] = []
    previous_space = False
    for original_index, char in enumerate(value):
        if char.isspace():
            if chars and not previous_space:
                chars.append(" ")
                mapping.append(original_index)
            previous_space = True
            continue
        chars.append(char)
        mapping.append(original_index)
        previous_space = False

    normalized = "".join(chars).strip()
    leading_trim = len(chars) - len("".join(chars).lstrip())
    if leading_trim:
        mapping = mapping[leading_trim:]
    trailing_len = len(normalized)
    mapping = mapping[:trailing_len]

    insertions = list(re.finditer(r"(?i)\bSr\.?\(a\)\.?(?=[A-Z\u00C0-\u00DE])", normalized))
    for match in reversed(insertions):
        insert_at = match.end()
        normalized = normalized[:insert_at] + " " + normalized[insert_at:]
        mapping.insert(insert_at, mapping[insert_at - 1] if insert_at > 0 and mapping else 0)
    return normalized, mapping


def _comma_parts_with_spans(text: str, base_start: int) -> list[dict[str, int | str]]:
    parts: list[dict[str, int | str]] = []
    for match in re.finditer(r"[^,]+", text):
        value = _clean_value(match.group(0))
        if not value:
            continue
        leading = len(match.group(0)) - len(match.group(0).lstrip())
        trailing = len(match.group(0).rstrip())
        start = base_start + match.start() + leading
        end = base_start + match.start() + trailing
        parts.append({"value": value, "start": start, "end": end})
    return parts


def _find_nationality(parts: list[dict[str, int | str]]) -> dict[str, int | str] | None:
    for part in parts:
        if _lower_ascii(str(part["value"])) in {_lower_ascii(item) for item in NATIONALITY_WORDS}:
            return part
    return parts[0] if parts else None


def _find_marital_status(parts: list[dict[str, int | str]]) -> dict[str, int | str] | None:
    for part in parts:
        if re.fullmatch(MARITAL_PATTERN, str(part["value"]), flags=re.IGNORECASE):
            return part
    for part in parts:
        match = re.search(MARITAL_PATTERN, str(part["value"]), flags=re.IGNORECASE)
        if match:
            return {
                "value": match.group(0),
                "start": int(part["start"]) + match.start(),
                "end": int(part["start"]) + match.end(),
            }
    return None


def _is_rejected_profession_part(value: str) -> bool:
    normalized = _lower_ascii(value)
    if not normalized or normalized in {"cpf", "comprador", "lote", "quadra"}:
        return True
    if normalized in {_lower_ascii(item) for item in NATIONALITY_WORDS}:
        return True
    return bool(re.fullmatch(MARITAL_PATTERN, value, flags=re.IGNORECASE))


def _global_span(section: Any, local_start: int, local_end: int) -> tuple[int, int]:
    return section.start_index + local_start, section.start_index + local_end


def _original_span(index_map: list[int], normalized_start: int, normalized_end: int) -> tuple[int, int]:
    if not index_map or normalized_start < 0 or normalized_end <= normalized_start:
        return -1, -1
    start = index_map[min(normalized_start, len(index_map) - 1)]
    end = index_map[min(normalized_end - 1, len(index_map) - 1)] + 1
    return start, end


def _clean_value(value: str) -> str:
    clean = " ".join((value or "").split())
    return clean.strip(" ,.;:-")


def _clean_name(value: str) -> str:
    clean = _clean_value(value)
    clean = re.sub(r"(?i)^\s*DECLARA(?:C|\u00C7)(?:AO|\u00C3O)?\s+", "", clean)
    clean = re.sub(r"(?i)^(?:Sr\.?\(a\)\.?\s*|Sra?\.?\s*|Senhor(?:a)?\s+)+", "", clean)
    clean = re.split(r"\b(?:PARA|QUALIFICAD[OA]|INSCRIT[OA]|PORTADOR(?:A)?|REPRESENTAD[OA])\b", clean, maxsplit=1)[0]
    return clean.strip(" ,.;:-")


def _clean_enterprise(value: str) -> str:
    clean = _clean_value(value)
    clean = re.split(r"\b(?:ADQUIRIDO|LOCALIZADO|SITUADO)\b", clean, maxsplit=1, flags=re.IGNORECASE)[0]
    return clean.strip(" ,.;:-")


def _clean_city(value: str) -> str:
    clean = _clean_value(value)
    tokens = clean.split()
    for index, token in enumerate(tokens):
        if any(char.islower() for char in token):
            return " ".join(tokens[index:]) if index else clean
    return clean


def _lower_ascii(value: str) -> str:
    decomposed = unicodedata.normalize("NFKD", value or "")
    ascii_text = "".join(char for char in decomposed if not unicodedata.combining(char))
    return ascii_text.lower().strip()


def _same_text(left: str, right: str) -> bool:
    return _lower_ascii(left) == _lower_ascii(right)


def _snippet(text: str, start: int, end: int, radius: int = 80) -> str:
    left = max(0, start - radius)
    right = min(len(text), end + radius)
    snippet = " ".join(text[left:right].split())
    if left:
        snippet = "..." + snippet
    if right < len(text):
        snippet += "..."
    return snippet


def _word_boundary_pattern(value: str) -> str:
    escaped = re.escape(_clean_value(value))
    return rf"(?<![0-9{LETTER_RANGE}]){escaped}(?![0-9{LETTER_RANGE}])"


def _collect_docx_text(path: Path) -> str:
    document = Document(path)
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text:
            blocks.append(paragraph.text)
    for table in document.tables:
        _collect_table_text(table, blocks)
    for part in DOCXWriter._iter_section_parts(document):
        for paragraph in part.paragraphs:
            if paragraph.text:
                blocks.append(paragraph.text)
        for table in part.tables:
            _collect_table_text(table, blocks)
    return "\n".join(blocks)


def _collect_table_text(table: Any, blocks: list[str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if paragraph.text:
                    blocks.append(paragraph.text)
            for nested_table in cell.tables:
                _collect_table_text(nested_table, blocks)


def _replace_in_document(document: Any, replacements: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        DOCXWriter._apply_replacements(paragraph, replacements)
    for table in document.tables:
        _replace_in_table(table, replacements)
    for part in DOCXWriter._iter_section_parts(document):
        for paragraph in part.paragraphs:
            DOCXWriter._apply_replacements(paragraph, replacements)
        for table in part.tables:
            _replace_in_table(table, replacements)


def _replace_in_table(table: Any, replacements: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                DOCXWriter._apply_replacements(paragraph, replacements)
            for nested_table in cell.tables:
                _replace_in_table(nested_table, replacements)
