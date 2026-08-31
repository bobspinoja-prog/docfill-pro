from __future__ import annotations

import hashlib
import re
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document

from services.field_extractor import MARKER_BY_FIELD, extract_fields
from services.runtime_json_store import RuntimeJsonStore


FIELD_LABELS = {
    "{{COMPRADOR}}": "Nome Completo",
    "{{NACIONALIDADE}}": "Nacionalidade",
    "{{PROFISSAO}}": "Profiss\u00e3o",
    "{{ESTADO_CIVIL}}": "Estado Civil",
    "{{CPF_CNPJ}}": "CPF/CNPJ",
    "{{LOTE}}": "Lote/Unidade",
    "{{QUADRA}}": "Quadra",
    "{{EMPREENDIMENTO}}": "Empreendimento",
    "{{VENDEDOR}}": "Nome do Vendedor",
    "{{CIDADE}}": "Cidade",
    "{{DATA}}": "Data",
}

SUPPORTED_FIELDS = tuple(FIELD_LABELS)


@dataclass
class SemanticDetection:
    field: str
    value: str
    confidence: float
    source: str
    snippet: str

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)

    @classmethod
    def from_dict(cls, data: dict[str, Any]) -> "SemanticDetection":
        return cls(
            field=str(data.get("field", "")),
            value=str(data.get("value", "")),
            confidence=float(data.get("confidence", 0.0)),
            source=str(data.get("source", "")),
            snippet=str(data.get("snippet", "")),
        )


class TemplateSemanticAnalyzer(RuntimeJsonStore):
    """Infers values for the fixed DocFill Pro UI fields from a DOCX template."""

    filename = "template_semantic_mappings.json"
    default_content: dict[str, Any] = {}

    def analyze(self, template_path: str | Path) -> dict[str, SemanticDetection]:
        path = Path(template_path)
        template_hash = self.template_hash(path)
        blocks = self._collect_blocks(path)
        full_text = "\n".join(blocks)

        auto_detections: dict[str, SemanticDetection] = {}
        self._detect_placeholders(full_text, auto_detections)
        self._detect_semantic_fields(full_text, auto_detections)

        saved = self.load_template_mapping(template_hash)
        accepted = saved.get("accepted", {}) if isinstance(saved, dict) else {}
        detections = dict(auto_detections)
        for field, data in accepted.items():
            if field in SUPPORTED_FIELDS and isinstance(data, dict):
                detection = SemanticDetection.from_dict(data)
                if detection.value:
                    detections[field] = detection

        self.save_detections(template_hash, path.name, auto_detections, detections)
        return detections

    def save_detections(
        self,
        template_hash: str,
        template_name: str,
        auto_detections: dict[str, SemanticDetection],
        detections: dict[str, SemanticDetection],
    ) -> None:
        data = self.load()
        current = data.get(template_hash, {})
        data[template_hash] = {
            **current,
            "template_name": template_name,
            "updated_at": datetime.now().isoformat(timespec="seconds"),
            "usage_count": int(current.get("usage_count", 0)) + 1,
            "last_used_at": datetime.now().isoformat(timespec="seconds"),
            "auto_detections": {field: detection.to_dict() for field, detection in auto_detections.items()},
            "detections": {field: detection.to_dict() for field, detection in detections.items()},
            "accepted": current.get("accepted", {}),
            "history": current.get("history", []),
        }
        self.save(data)

    def accept_detection(self, template_hash: str, detection: SemanticDetection) -> None:
        data = self.load()
        item = data.setdefault(template_hash, {})
        accepted = item.setdefault("accepted", {})
        accepted[detection.field] = detection.to_dict()
        history = item.setdefault("history", [])
        history.append(
            {
                "at": datetime.now().isoformat(timespec="seconds"),
                "action": "accept",
                "field": detection.field,
                "value": detection.value,
                "confidence": detection.confidence,
                "source": detection.source,
            }
        )
        item["updated_at"] = datetime.now().isoformat(timespec="seconds")
        self.save(data)

    def save_manual_value(self, template_hash: str, field: str, value: str) -> SemanticDetection:
        detection = SemanticDetection(
            field=field,
            value=value.strip(),
            confidence=1.0,
            source="manual: correcao do usuario",
            snippet=value.strip(),
        )
        self.accept_detection(template_hash, detection)
        return detection

    def load_template_mapping(self, template_hash: str) -> dict[str, Any]:
        return dict(self.load().get(template_hash, {}))

    @staticmethod
    def template_hash(template_path: str | Path) -> str:
        handle = hashlib.sha256()
        with Path(template_path).open("rb") as file:
            for chunk in iter(lambda: file.read(1024 * 1024), b""):
                handle.update(chunk)
        return handle.hexdigest()

    @staticmethod
    def suggestion_values(detections: dict[str, SemanticDetection]) -> dict[str, str]:
        return {
            field: detection.value
            for field, detection in detections.items()
            if detection.value and field in SUPPORTED_FIELDS
        }

    def _detect_placeholders(self, text: str, detections: dict[str, SemanticDetection]) -> None:
        for match in re.finditer(r"\{\{[^{}]+\}\}", text):
            marker = self._normalize_marker(match.group(0))
            if marker not in SUPPORTED_FIELDS or marker in detections:
                continue
            detections[marker] = SemanticDetection(
                field=marker,
                value="",
                confidence=0.99,
                source="placeholder",
                snippet=self._snippet(text, match.start(), match.end()),
            )

    def _detect_semantic_fields(self, text: str, detections: dict[str, SemanticDetection]) -> None:
        extraction = extract_fields(text)
        for field_name, item in extraction.fields.items():
            marker = MARKER_BY_FIELD.get(field_name)
            if marker not in SUPPORTED_FIELDS or not item.value:
                continue
            self._add_detection(
                detections,
                marker,
                item.value,
                item.confidence,
                item.source,
                item.evidence,
            )

    def _add_detection(
        self,
        detections: dict[str, SemanticDetection],
        field: str,
        value: str,
        confidence: float,
        source: str,
        snippet: str,
    ) -> None:
        if field not in SUPPORTED_FIELDS:
            return
        clean_value = self._clean_value(value)
        if not clean_value:
            return
        current = detections.get(field)
        if current and current.value and current.confidence >= confidence:
            return
        detections[field] = SemanticDetection(
            field=field,
            value=clean_value,
            confidence=round(min(1.0, max(0.0, confidence)), 2),
            source=source,
            snippet=snippet,
        )

    def _collect_blocks(self, template_path: Path) -> list[str]:
        document = Document(template_path)
        blocks: list[str] = []
        for paragraph in document.paragraphs:
            self._append_block(blocks, paragraph.text)
        for table in document.tables:
            self._append_table(blocks, table)
        for part in self._iter_section_parts(document):
            for paragraph in part.paragraphs:
                self._append_block(blocks, paragraph.text)
            for table in part.tables:
                self._append_table(blocks, table)
        return blocks

    @classmethod
    def _append_table(cls, blocks: list[str], table: Any) -> None:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    cls._append_block(blocks, paragraph.text)
                for nested_table in cell.tables:
                    cls._append_table(blocks, nested_table)

    @staticmethod
    def _append_block(blocks: list[str], text: str) -> None:
        clean = " ".join((text or "").split())
        if clean:
            blocks.append(clean)

    @staticmethod
    def _iter_section_parts(document: Any) -> Any:
        seen: set[int] = set()
        attrs = (
            "header",
            "first_page_header",
            "even_page_header",
            "footer",
            "first_page_footer",
            "even_page_footer",
        )
        for section in document.sections:
            for attr_name in attrs:
                part = getattr(section, attr_name, None)
                if part is None:
                    continue
                element_id = id(part._element)
                if element_id in seen:
                    continue
                seen.add(element_id)
                yield part

    @staticmethod
    def _snippet(text: str, start: int, end: int, radius: int = 70) -> str:
        left = max(0, start - radius)
        right = min(len(text), end + radius)
        snippet = " ".join(text[left:right].split())
        if left > 0:
            snippet = "..." + snippet
        if right < len(text):
            snippet += "..."
        return snippet

    @staticmethod
    def _clean_value(value: str) -> str:
        clean = " ".join((value or "").split())
        clean = clean.strip(" ,.;:-")
        return clean

    @staticmethod
    def _normalize_marker(marker: str) -> str:
        marker_text = marker.strip().upper()
        marker_text = marker_text.removeprefix("{{").removesuffix("}}").strip()
        marker_text = marker_text.strip("{}").strip()
        return f"{{{{{marker_text}}}}}" if marker_text else ""

