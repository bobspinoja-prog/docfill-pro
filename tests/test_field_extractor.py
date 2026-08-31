from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from services.docx_reader import DOCXReader
from services.field_extractor import FIELDS, extract_fields, normalize_text, rewrite_template_with_markers
from services.text_sections import split_text_sections


REAL_PREAMBLE_TEXT = (
    "STEFANO DO COUTO ROSA MILO, brasileiro, advogado, solteiro, administrador "
    "portador do CPF n\u00b0 436.106.638-80, na qualidade de COMPRADOR do Lote 15 "
    "Quadra 13A do LOTEAMENTO ALPHAVILLE RIBEIR\u00c3O PRETO, adquirido atrav\u00e9s de "
    "Instrumento Particular de Promessa de Venda e Compra do(a) Sr.(a).CARLOS "
    "ALBERTO CHAIN CAMPANA, para os devidos fins de direito e sob as penas da Lei..."
)

DATE_TEXT = "Ribeir\u00e3o Preto, 17 de Junho de 2026"
SIGNATURES = "CARLOS ALBERTO CHAIN CAMPANA    STEFANO DO COUTO ROSA MILO"


def test_normalize_text_fixes_spacing_and_joins_lines() -> None:
    normalized = normalize_text("Sr.(a).CARLOS\n\n  ALBERTO")

    assert normalized == "Sr.(a). CARLOS ALBERTO"


def test_text_sections_split_required_blocks() -> None:
    text = (
        f"{REAL_PREAMBLE_TEXT} A) RECEBI o valor. "
        "Assim, por todo o exposto, declaro. "
        f"{DATE_TEXT} {SIGNATURES}"
    )
    sections = split_text_sections(normalize_text(text))

    assert "STEFANO DO COUTO ROSA MILO" in sections.get("preambulo").text
    assert sections.get("corpo").text.startswith("A) RECEBI")
    assert sections.get("paragrafo_final").text.startswith("Assim, por todo o exposto")
    assert sections.get("data").text == DATE_TEXT
    assert "CARLOS ALBERTO CHAIN CAMPANA" in sections.get("assinaturas").text


def test_extract_fields_from_required_real_preamble() -> None:
    extraction = extract_fields(f"{REAL_PREAMBLE_TEXT} {DATE_TEXT} {SIGNATURES}")
    values = extraction.as_plain_values(min_confidence=0.60)

    assert values["COMPRADOR"] == "STEFANO DO COUTO ROSA MILO"
    assert values["NACIONALIDADE"] == "brasileiro"
    assert values["PROFISSAO"] == "advogado / administrador"
    assert values["ESTADO_CIVIL"] == "solteiro"
    assert values["CPF_CNPJ"] == "436.106.638-80"
    assert values["LOTE"] == "15"
    assert values["QUADRA"] == "13A"
    assert values["EMPREENDIMENTO"] == "LOTEAMENTO ALPHAVILLE RIBEIR\u00c3O PRETO"
    assert values["VENDEDOR"] == "CARLOS ALBERTO CHAIN CAMPANA"


def test_extract_fields_from_required_date_line() -> None:
    extraction = extract_fields(DATE_TEXT)
    values = extraction.as_plain_values(min_confidence=0.60)

    assert values["CIDADE"] == "Ribeir\u00e3o Preto"
    assert values["DATA"] == "17 de Junho de 2026"


def test_extract_fields_returns_location_confidence_and_occurrences() -> None:
    text = f"{REAL_PREAMBLE_TEXT} {DATE_TEXT} {SIGNATURES}"
    result = extract_fields(text).to_dict()

    assert tuple(result) == FIELDS
    for field, payload in result.items():
        assert {
            "value",
            "confidence",
            "source",
            "reason",
            "start_index",
            "end_index",
            "section",
            "marker",
            "occurrences",
        }.issubset(payload)
        if payload["value"]:
            assert payload["confidence"] >= 0.60, field
            assert payload["source"] != "not_found", field
            assert payload["start_index"] >= 0, field
            assert payload["end_index"] > payload["start_index"], field
            assert text[payload["start_index"]:payload["end_index"]] in text

    profissao = result["PROFISSAO"]
    assert profissao["value"] == "advogado / administrador"
    assert [item["value"] for item in profissao["occurrences"]] == ["advogado", "administrador"]
    assert result["COMPRADOR"]["marker"] == "{{COMPRADOR}}"
    assert result["VENDEDOR"]["section"] == "preambulo"


def test_seller_conflict_prefers_acquisition_clause_and_marks_conflict_no_replace() -> None:
    text = (
        f"{REAL_PREAMBLE_TEXT} Assim, por todo o exposto, responsabilidade sobre "
        "o meu compromisso firmado com o(a) Sr.(a). JOAO FINAL ERRADO."
    )
    result = extract_fields(text).to_dict()

    assert result["VENDEDOR"]["value"] == "CARLOS ALBERTO CHAIN CAMPANA"
    assert result["VENDEDOR"]["confidence"] < 0.97
    conflict = result["VENDEDOR"]["occurrences"][-1]
    assert conflict["value"] == "JOAO FINAL ERRADO"
    assert conflict["section"] == "paragrafo_final"
    assert conflict["replace"] is False


def test_final_paragraph_seller_is_low_confidence_fallback() -> None:
    text = "Assim, por todo o exposto, responsabilidade sobre o meu compromisso firmado com o(a) Sr.(a). JOAO FINAL ERRADO."
    result = extract_fields(text).to_dict()

    assert result["VENDEDOR"]["value"] == "JOAO FINAL ERRADO"
    assert result["VENDEDOR"]["confidence"] == 0.55
    assert result["VENDEDOR"]["source"] == "fallback"


def test_cnpj_lote_with_zero_and_quadra_with_letter() -> None:
    text = (
        "EMPRESA CLIENTE TESTE LTDA, brasileira, comerciante, casada portadora do "
        "CNPJ n\u00ba 62.990.091/0001-07, na qualidade de COMPRADOR do Lote 04 Quadra 08B "
        "do CONDOMINIO ALPHAVILLE TESTE, adquirido atraves de Instrumento Particular "
        "de Compra e Venda do(a) Sr.(a).VENDEDOR TESTE FINAL, para fins."
    )
    values = extract_fields(text).as_plain_values(min_confidence=0.60)

    assert values["CPF_CNPJ"] == "62.990.091/0001-07"
    assert values["LOTE"] == "04"
    assert values["QUADRA"] == "08B"
    assert values["EMPREENDIMENTO"] == "CONDOMINIO ALPHAVILLE TESTE"


def test_ocr_degraded_accents_with_question_marks_still_extracts_preamble() -> None:
    text = (
        "STEFANO DO COUTO ROSA MILO, brasileiro, advogado, solteiro, administrador "
        "portador do CPF n? 436.106.638-80, na qualidade de COMPRADOR do Lote 15 "
        "Quadra 13A do LOTEAMENTO ALPHAVILLE RIBEIR?O PRETO, adquirido atrav?s de "
        "Instrumento Particular de Promessa de Venda e Compra do(a) Sr.(a).CARLOS "
        "ALBERTO CHAIN CAMPANA, para os devidos fins. Ribeir?o Preto, 17 de Junho de 2026"
    )
    values = extract_fields(text).as_plain_values(min_confidence=0.60)

    assert values["COMPRADOR"] == "STEFANO DO COUTO ROSA MILO"
    assert values["CPF_CNPJ"] == "436.106.638-80"
    assert values["EMPREENDIMENTO"] == "LOTEAMENTO ALPHAVILLE RIBEIR?O PRETO"
    assert values["VENDEDOR"] == "CARLOS ALBERTO CHAIN CAMPANA"


def test_pdf_title_prefix_is_not_part_of_buyer_name() -> None:
    text = (
        "DECLARAÇÃO RENAN SARTOR PIESTCH, brasileiro, empresário, casado, "
        "portador do CPF n° 058.8836.619-63 na qualidade de COMPRADOR do Lote 20 "
        "Quadra 5A do LOTEAMENTO ALPHAVILLE RIBEIRÃO PRETO, vendido atraves do "
        "contrato de compra e venda de sobre imóvel com permuta do(a) ROBERTO "
        "PADUA VALADAO JUNIOR para os devidos fins. Ribeirão Preto, 23 de Junho de 2026"
    )
    values = extract_fields(text).as_plain_values(min_confidence=0.60)

    assert values["COMPRADOR"] == "RENAN SARTOR PIESTCH"


def test_city_line_drops_signature_name_before_pdf_date() -> None:
    text = (
        "Assim, por todo o exposto, firmado com o(a) Sr.(a). ROBERTO PADUA VALADAO "
        "Ribeirão Preto, 23 de Junho de 2026 ________________________________________ "
        "ROBERTO PADUA VALADAO RENAN SARTOR PIESTCH"
    )
    values = extract_fields(text).as_plain_values(min_confidence=0.60)

    assert values["CIDADE"] == "Ribeirão Preto"


def test_document_without_field_keeps_empty_payload() -> None:
    result = extract_fields("Documento sem preambulo estruturado.").to_dict()

    assert result["COMPRADOR"]["value"] == ""
    assert result["COMPRADOR"]["confidence"] == 0.0
    assert result["COMPRADOR"]["start_index"] == -1
    assert result["COMPRADOR"]["occurrences"] == []


def test_rewrite_template_with_markers_preserves_original_and_marks_safe_values() -> None:
    with TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        input_docx = tmp_path / "declaracao.docx"
        output_docx = tmp_path / "TEMPLATE_MARCADO - declaracao.docx"

        doc = Document()
        doc.add_paragraph(REAL_PREAMBLE_TEXT)
        doc.add_paragraph(DATE_TEXT)
        doc.save(input_docx)

        extraction = extract_fields(f"{REAL_PREAMBLE_TEXT} {DATE_TEXT}")
        report = rewrite_template_with_markers(input_docx, output_docx, extraction)

        original_text = DOCXReader(input_docx).extract_text({})
        marked_text = DOCXReader(output_docx).extract_text({})

        assert "STEFANO DO COUTO ROSA MILO" in original_text
        assert "{{COMPRADOR}}" in marked_text
        assert "{{NACIONALIDADE}}" in marked_text
        assert "{{PROFISSAO}}" in marked_text
        assert "{{ESTADO_CIVIL}}" in marked_text
        assert "{{CPF_CNPJ}}" in marked_text
        assert "{{LOTE}}" in marked_text
        assert "{{QUADRA}}" in marked_text
        assert "{{EMPREENDIMENTO}}" in marked_text
        assert "{{VENDEDOR}}" in marked_text
        assert "{{CIDADE}}" in marked_text
        assert "{{DATA}}" in marked_text
        assert report.output_path == output_docx
        assert "COMPRADOR" in report.marked_fields
