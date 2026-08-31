from __future__ import annotations

import hashlib
import sys
from datetime import datetime
from pathlib import Path
from tempfile import TemporaryDirectory

import customtkinter as ctk
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from services.docx_reader import DOCXReader
from services.docx_writer import DOCXWriter
from services.history_manager import HistoryManager
from services.history_suggestions import HistorySuggestion, HistorySuggestions
from services.mapping_manager import MappingManager
from services.structured_logger import StructuredLogger
from services.template_profile_store import TemplateProfileStore
from services.semantic_replacements import build_safe_semantic_replacements
from services.template_semantic_analyzer import SemanticDetection, TemplateSemanticAnalyzer
from services.user_session_store import UserSessionStore
from ui.main_window import DocFillProApp


def file_hash(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def test_marker_docx_roundtrip() -> None:
    with TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        template = tmp_path / "template.docx"
        output = tmp_path / "output.docx"

        doc = Document()
        doc.add_paragraph("Comprador: {{COMPRADOR}}")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "CPF"
        table.cell(0, 1).text = "{{CPF_CNPJ}}"
        section = doc.sections[0]
        section.header.paragraphs[0].text = "Cidade {{CIDADE}}"
        section.footer.paragraphs[0].text = "Data {{DATA}}"
        doc.save(template)

        replacements = {
            "{{COMPRADOR}}": "Cliente Teste",
            "{{CPF_CNPJ}}": "111.222.333-44",
            "{{CIDADE}}": "Campinas",
            "{{DATA}}": "12/06/2026",
        }

        before = file_hash(template)
        DOCXWriter().generate(template, output, replacements)
        assert before == file_hash(template), "template original foi alterado"

        generated_text = DOCXReader(output).extract_text({})
        for value in replacements.values():
            assert value in generated_text
        assert "{{" not in generated_text


def test_marker_split_across_runs_is_replaced() -> None:
    with TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        template = tmp_path / "split.docx"
        output = tmp_path / "split_output.docx"

        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("Olá ")
        paragraph.add_run("{{COMPRADOR}}")
        paragraph.add_run("!")
        doc.save(template)

        DOCXWriter().generate(template, output, {"{{COMPRADOR}}": "Cliente Teste"})

        generated_text = DOCXReader(output).extract_text({})

        assert "Cliente Teste" in generated_text
        assert "{{COMPRADOR}}" not in generated_text


def test_repeated_marker_is_fully_replaced() -> None:
    with TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        template = tmp_path / "repeated.docx"
        output = tmp_path / "repeated_output.docx"

        doc = Document()
        doc.add_paragraph(" ".join(["{{COMPRADOR}}"] * 12))
        doc.save(template)

        DOCXWriter().generate(template, output, {"{{COMPRADOR}}": "Cliente"})

        generated_text = DOCXReader(output).extract_text({})

        assert generated_text.count("Cliente") == 12
        assert "{{COMPRADOR}}" not in generated_text


def test_mapping_priority_and_filename() -> None:
    with TemporaryDirectory() as tmp:
        mapping = MappingManager(Path(tmp) / "mappings.json")
        mapping.add_marker("{{COMPRADOR}}", "Custom")
        replacements = mapping.build_replacements({"{{COMPRADOR}}": "Principal"})
        assert replacements["{{COMPRADOR}}"] == "Principal"

    safe_name = DocFillProApp._sanitize_filename('Maria: Silva / Teste?* .')
    assert safe_name
    assert not any(char in safe_name for char in '<>:"/\\|?*')
    assert not safe_name.endswith((".", " "))


def test_ui_layout() -> None:
    ctk.set_appearance_mode("dark")
    ctk.set_default_color_theme("green")
    app = DocFillProApp()
    app.geometry("1366x768")
    app.update()
    assert len(app.form_panel.entries) == 11
    assert app.form_panel.winfo_width() >= 420
    assert app.preview_panel.winfo_width() >= 600
    app.close_app()


def test_history_suggestion_widget_toggles_visibility() -> None:
    ctk.set_appearance_mode("dark")
    ctk.set_default_color_theme("green")
    app = DocFillProApp()
    suggestion = HistorySuggestion(
        field="{{CPF_CNPJ}}",
        value="123.456.789-00",
        source="history",
        confidence=0.92,
        reason="Mesmo comprador encontrado em histórico anterior",
        current_value="",
        status="empty",
        anchor_field="{{COMPRADOR}}",
        anchor_value="EDUARDO TEIXEIRA",
        template_hash="hash-1",
        template_name="modelo.docx",
        record_id="record-1",
        key="hash-1|{{COMPRADOR}}|eduardo teixeira|{{CPF_CNPJ}}|12345678900",
    )
    app.form_panel.set_history_suggestion(
        "{{CPF_CNPJ}}",
        suggestion,
        apply_command=lambda: None,
        ignore_command=lambda: None,
    )
    app.update()
    assert app.form_panel.suggestion_labels["{{CPF_CNPJ}}"].cget("text").startswith("Sugestão do histórico")
    assert app.form_panel.suggestion_frames["{{CPF_CNPJ}}"].winfo_ismapped()
    app.form_panel.clear_history_suggestion("{{CPF_CNPJ}}")
    app.update()
    assert not app.form_panel.suggestion_frames["{{CPF_CNPJ}}"].winfo_ismapped()
    app.close_app()


def test_example_document_if_available() -> None:
    downloads = Path.home() / "Downloads"
    examples = [
        path
        for path in downloads.glob("*.docx")
        if "exemplo app" in path.name and not path.name.startswith("~$")
    ]
    if not examples:
        print("example docx not found; skipping example-specific smoke test")
        return

    template = examples[0]
    reader = DOCXReader(template)
    suggestions = reader.suggest_values()
    if not {"{{COMPRADOR}}", "{{CPF_CNPJ}}", "{{VENDEDOR}}"}.issubset(suggestions):
        print(f"example docx {template.name} is not the expected fixture; skipping example-specific smoke test")
        return

    assert suggestions["{{COMPRADOR}}"].startswith("ANDERSON")
    assert suggestions["{{CPF_CNPJ}}"] == "269.199.688-35"
    assert suggestions["{{VENDEDOR}}"].startswith("RODRIGO")

    values = {
        "{{COMPRADOR}}": "CLIENTE TESTE DOCFILL",
        "{{NACIONALIDADE}}": "brasileira",
        "{{PROFISSAO}}": "engenheira",
        "{{ESTADO_CIVIL}}": "solteira",
        "{{CPF_CNPJ}}": "111.222.333-44",
        "{{LOTE}}": "99",
        "{{QUADRA}}": "10A",
        "{{EMPREENDIMENTO}}": "LOTEAMENTO TESTE VERDE",
        "{{VENDEDOR}}": "VENDEDOR TESTE DOCFILL",
        "{{CIDADE}}": "Campinas",
        "{{DATA}}": "12 de Junho de 2026",
    }
    replacements = MappingManager(ROOT / "data" / "mappings.json").build_replacements(values)
    replacements.update(reader.build_literal_replacements(values))

    output_dir = ROOT / "test_outputs"
    output_dir.mkdir(exist_ok=True)
    output = output_dir / "SMOKE_TEST_EXEMPLO.docx"
    if output.exists():
        output.unlink()

    before = file_hash(template)
    DOCXWriter().generate(template, output, replacements)
    assert before == file_hash(template), "template exemplo foi alterado"

    generated_text = DOCXReader(output).extract_text({})
    for expected in (
        "CLIENTE TESTE DOCFILL",
        "111.222.333-44",
        "Lote 99",
        "Quadra 10A",
        "LOTEAMENTO TESTE VERDE",
        "VENDEDOR TESTE DOCFILL",
        "Campinas, 12 de Junho de 2026",
    ):
        assert expected in generated_text


def test_semantic_template_detection_and_persistence() -> None:
    with TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        template = tmp_path / "semantic.docx"
        mapping_file = tmp_path / "template_semantic_mappings.json"

        doc = Document()
        doc.add_paragraph(
            "permituta do(a) EDUARDO TEIXEIRA, brasileira, engenheira, solteira, portadora do CNPJ nº 62.990.091/0001-07 "
            "na qualidade de COMPRADOR do Lote 04 Quadra 08B do LOTEAMENTO ALPHAVILLE RIBEIRÃO PRETO, adquirido do(a) "
            "RODRIGO SILVA, para fins."
        )
        doc.add_paragraph("RIBEIRÃO PRETO, 15 de Junho de 2026")
        doc.save(template)

        analyzer = TemplateSemanticAnalyzer(mapping_file)
        detections = analyzer.analyze(template)

        assert detections["{{COMPRADOR}}"].value == "EDUARDO TEIXEIRA"
        assert detections["{{CPF_CNPJ}}"].value == "62.990.091/0001-07"
        assert detections["{{LOTE}}"].value == "04"
        assert detections["{{QUADRA}}"].value == "08B"
        assert detections["{{EMPREENDIMENTO}}"].value == "LOTEAMENTO ALPHAVILLE RIBEIRÃO PRETO"
        assert detections["{{VENDEDOR}}"].value == "RODRIGO SILVA"
        assert detections["{{CIDADE}}"].value == "RIBEIRÃO PRETO"
        assert detections["{{DATA}}"].value == "15 de Junho de 2026"

        template_hash = analyzer.template_hash(template)
        analyzer.save_manual_value(template_hash, "{{COMPRADOR}}", "CLIENTE CORRIGIDO")
        detections_after = analyzer.analyze(template)
        assert detections_after["{{COMPRADOR}}"].value == "CLIENTE CORRIGIDO"
        saved = analyzer.load_template_mapping(template_hash)
        assert "auto_detections" in saved
        assert saved["usage_count"] >= 2


def test_safe_semantic_replacements_block_ambiguous_name() -> None:
    text = "RIBEIRAO PRETO apareceu no contrato. Depois RIBEIRAO PRETO confirmou. RIBEIRAO PRETO nao deve ser afetado."
    result = build_safe_semantic_replacements(
        {
            "{{CIDADE}}": {
                "field": "{{CIDADE}}",
                "value": "RIBEIRAO PRETO",
                "confidence": 0.95,
                "source": "context: cidade",
                "snippet": "RIBEIRAO PRETO",
            }
        },
        {"{{CIDADE}}": "CAMPINAS"},
        {},
        text,
    )
    assert result.safe_replacements == {}
    assert result.blocked_replacements
    assert any("múltiplos contextos" in warning for warning in result.warnings)


def test_safe_semantic_replacements_allow_unique_cpf_and_dates() -> None:
    text = "CPF 62.990.091/0001-07 e data 15 de Junho de 2026."
    result = build_safe_semantic_replacements(
        {
            "{{CPF_CNPJ}}": {
                "field": "{{CPF_CNPJ}}",
                "value": "62.990.091/0001-07",
                "confidence": 0.96,
                "source": "context: CPF/CNPJ",
                "snippet": "62.990.091/0001-07",
            },
            "{{DATA}}": {
                "field": "{{DATA}}",
                "value": "15 de Junho de 2026",
                "confidence": 0.92,
                "source": "context: data",
                "snippet": "15 de Junho de 2026",
            },
        },
        {"{{CPF_CNPJ}}": "111.222.333-44", "{{DATA}}": "12 de Junho de 2026"},
        {},
        text,
    )
    assert result.safe_replacements["62.990.091/0001-07"] == "111.222.333-44"
    assert result.safe_replacements["15 de Junho de 2026"] == "12 de Junho de 2026"


def test_mapping_seed_and_build_files_are_present() -> None:
    seed = ROOT / "data" / "template_semantic_mappings.json"
    assert seed.exists()
    assert seed.read_text(encoding="utf-8").strip() == "{}"
    assert "template_semantic_mappings.json" in (ROOT / "DOCFILL PRO.spec").read_text(encoding="utf-8")
    assert "template_semantic_mappings.json" in (ROOT / "installer" / "DOCFILL_PRO_Inno.iss").read_text(encoding="utf-8")
    assert (ROOT / "data" / "template_profiles.json").exists()
    assert (ROOT / "data" / "history.json").exists()
    assert (ROOT / "data" / "user_session.json").exists()
    assert "template_profiles.json" in (ROOT / "DOCFILL PRO.spec").read_text(encoding="utf-8")
    assert "history.json" in (ROOT / "DOCFILL PRO.spec").read_text(encoding="utf-8")
    assert "user_session.json" in (ROOT / "DOCFILL PRO.spec").read_text(encoding="utf-8")
    assert "template_profiles.json" in (ROOT / "installer" / "DOCFILL_PRO_Inno.iss").read_text(encoding="utf-8")
    assert "history.json" in (ROOT / "installer" / "DOCFILL_PRO_Inno.iss").read_text(encoding="utf-8")
    assert "user_session.json" in (ROOT / "installer" / "DOCFILL_PRO_Inno.iss").read_text(encoding="utf-8")


def test_history_manager_records_and_filters_documents() -> None:
    with TemporaryDirectory() as tmp:
        store = HistoryManager(Path(tmp) / "history.json")
        template_path = Path(tmp) / "template.docx"
        output_path = Path(tmp) / "saida" / "documento.docx"
        record = store.record_document(
            template_name=template_path.name,
            template_hash="hash-321",
            output_file=output_path,
            document_name=output_path.name,
            fields={"{{COMPRADOR}}": "EDUARDO TEIXEIRA"},
            detected_fields={
                "{{COMPRADOR}}": {
                    "field": "{{COMPRADOR}}",
                    "value": "EDUARDO TEIXEIRA",
                    "confidence": 0.88,
                    "source": "context: permituta do(a)",
                    "snippet": "permituta do(a) EDUARDO TEIXEIRA",
                }
            },
            profile_used="hash-321",
            template_path=template_path,
            output_folder=output_path.parent,
        )

        assert record["document_name"] == output_path.name
        assert store.get_record(record["id"]) is not None
        today = datetime.now().date().isoformat()
        filtered = store.query_records(search="EDUARDO", date_from=today, date_to=today)
        assert filtered
        assert filtered[0]["id"] == record["id"]
        assert store.set_favorite(record["id"], True) is not None
        favorites = store.query_records(favorites_only=True)
        assert favorites and favorites[0]["favorite"] is True
        assert store.recent_templates()[0]["template_hash"] == "hash-321"
        assert store.recent_documents()[0]["document_name"] == output_path.name


def test_history_suggestions_cover_matching_and_conflicts() -> None:
    with TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        history_store = HistoryManager(tmp_path / "history.json")
        session_store = UserSessionStore(tmp_path / "user_session.json")
        semantic_store = TemplateSemanticAnalyzer(tmp_path / "template_semantic_mappings.json")
        service = HistorySuggestions(history_store, session_store, semantic_store)

        assert service.build_suggestions({"{{COMPRADOR}}": "EDUARDO TEIXEIRA"}) == {}

        history_store.record_document(
            template_name="modelo.docx",
            template_hash="hash-aaa",
            output_file=tmp_path / "saida1.docx",
            document_name="saida1.docx",
            fields={
                "{{COMPRADOR}}": "ÉDUARDO TEIXEIRA",
                "{{CPF_CNPJ}}": "123.456.789-00",
                "{{NACIONALIDADE}}": "brasileira",
                "{{PROFISSAO}}": "engenheiro",
                "{{ESTADO_CIVIL}}": "solteiro",
                "{{CIDADE}}": "Ribeirão Preto",
                "{{VENDEDOR}}": "RODRIGO SILVA",
                "{{EMPREENDIMENTO}}": "ALPHAVILLE",
            },
            detected_fields={},
            profile_used="hash-aaa",
            template_path=tmp_path / "modelo.docx",
            output_folder=tmp_path / "saida",
        )

        suggestions = service.build_suggestions({"{{COMPRADOR}}": "eduardo teixeira"}, template_hash="hash-aaa")
        assert suggestions["{{CPF_CNPJ}}"].value == "123.456.789-00"
        assert suggestions["{{CPF_CNPJ}}"].status == "empty"
        assert suggestions["{{CPF_CNPJ}}"].confidence >= 0.72
        assert suggestions["{{CIDADE}}"].value == "Ribeirão Preto"

        ignored = service.build_suggestions(
            {"{{COMPRADOR}}": "eduardo teixeira"},
            template_hash="hash-aaa",
            ignored_keys={suggestions["{{CPF_CNPJ}}"].key},
        )
        assert "{{CPF_CNPJ}}" not in ignored

        history_store.record_document(
            template_name="modelo.docx",
            template_hash="hash-bbb",
            output_file=tmp_path / "saida2.docx",
            document_name="saida2.docx",
            fields={
                "{{COMPRADOR}}": "EDUARDO TEIXEIRA",
                "{{CPF_CNPJ}}": "111.111.111-11",
            },
            detected_fields={},
            profile_used="hash-bbb",
            template_path=tmp_path / "modelo.docx",
            output_folder=tmp_path / "saida",
        )

        conflicting = service.build_suggestions({"{{COMPRADOR}}": "EDUARDO TEIXEIRA"}, template_hash="hash-aaa")
        assert "{{CPF_CNPJ}}" not in conflicting

        service_same_template = HistorySuggestions(
            HistoryManager(tmp_path / "history_empty.json"),
            UserSessionStore(tmp_path / "user_session_empty.json"),
            TemplateSemanticAnalyzer(tmp_path / "template_semantic_empty.json"),
        )
        assert service_same_template.build_suggestions({"{{COMPRADOR}}": ""}) == {}


def test_template_profile_store_tracks_corrections() -> None:
    with TemporaryDirectory() as tmp:
        store = TemplateProfileStore(Path(tmp) / "template_profiles.json")
        detection = SemanticDetection(
            field="{{COMPRADOR}}",
            value="EDUARDO TEIXEIRA",
            confidence=0.93,
            source="context: permituta do(a)",
            snippet="permituta do(a) EDUARDO TEIXEIRA",
        )
        profile = store.update_profile(
            "hash-xyz",
            "modelo.docx",
            detections={detection.field: detection},
            placeholders=["{{COMPRADOR}}"],
            required_fields={"{{COMPRADOR}}"},
            manual_values={"{{COMPRADOR}}": "EDUARDO TEIXEIRA"},
        )
        assert profile["usage_count"] == 1
        assert "{{COMPRADOR}}" in profile["learned_fields"]

        store.record_correction("hash-xyz", "modelo.docx", detection)
        summary = store.summarize("hash-xyz")
        assert summary["corrections"]
        assert summary["template_name"] == "modelo.docx"


def test_structured_logger_writes_jsonl() -> None:
    with TemporaryDirectory() as tmp:
        logger = StructuredLogger(Path(tmp))
        log_path = logger.log("template_loaded", message="Template carregado", template_path="modelo.docx")
        assert log_path.exists()
        content = log_path.read_text(encoding="utf-8").strip().splitlines()
        assert content
        assert "template_loaded" in content[0]


def main() -> None:
    tests = [
        test_marker_docx_roundtrip,
        test_marker_split_across_runs_is_replaced,
        test_repeated_marker_is_fully_replaced,
        test_mapping_priority_and_filename,
        test_safe_semantic_replacements_block_ambiguous_name,
        test_safe_semantic_replacements_allow_unique_cpf_and_dates,
        test_mapping_seed_and_build_files_are_present,
        test_history_manager_records_and_filters_documents,
        test_history_suggestions_cover_matching_and_conflicts,
        test_template_profile_store_tracks_corrections,
        test_structured_logger_writes_jsonl,
        test_ui_layout,
        test_example_document_if_available,
        test_semantic_template_detection_and_persistence,
    ]
    for test in tests:
        test()
        print(f"ok - {test.__name__}")


if __name__ == "__main__":
    main()
